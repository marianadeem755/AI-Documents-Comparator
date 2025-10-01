import streamlit as st
import google.generativeai as genai
import openai
import PyPDF2
import io
import docx
from io import BytesIO
import os
from dotenv import load_dotenv
import difflib
import re
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import inch
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from reportlab.pdfgen import canvas
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import json
from datetime import datetime
import fitz  # pymupdf
from PIL import Image
import pytesseract
from langchain_core.prompts import PromptTemplate
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from typing import List
import uuid
import pandas as pd

# Load environment variables
load_dotenv()

# Configure page
st.set_page_config(
    page_title="AI Document Comparator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for attractive interface with keyword highlighting
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(45deg, #FF1744, #E91E63, #9C27B0, #673AB7);
        background-size: 300% 300%;
        animation: headerGlow 4s ease-in-out infinite;
        padding: 4rem 2rem;
        border-radius: 25px;
        margin-bottom: 3rem;
        text-align: center;
        color: white;
        box-shadow: 
            0 0 40px rgba(255, 23, 68, 0.4),
            0 20px 60px rgba(156, 39, 176, 0.3),
            inset 0 1px 0 rgba(255, 255, 255, 0.2);
        border: 4px solid transparent;
        background-clip: padding-box;
        position: relative;
        overflow: hidden;
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: -2px;
        left: -2px;
        right: -2px;
        bottom: -2px;
        background: linear-gradient(45deg, #FFD700, #FF69B4, #00BFFF, #32CD32);
        background-size: 400% 400%;
        animation: borderGlow 3s ease-in-out infinite;
        border-radius: 27px;
        z-index: -1;
    }
    
    @keyframes headerGlow {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    @keyframes borderGlow {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    .main-header h1 {
        font-size: 3.2rem;
        font-weight: 900;
        margin: 0 0 0.8rem 0;
        text-shadow: 
            2px 2px 8px rgba(0,0,0,0.6),
            0 0 20px rgba(255, 255, 255, 0.3),
            0 0 40px rgba(255, 23, 68, 0.5);
        color: #FFFFFF;
        letter-spacing: 2px;
        text-transform: uppercase;
        animation: textPulse 2s ease-in-out infinite;
    }
    
    @keyframes textPulse {
        0%, 100% { transform: scale(1); }
        50% { transform: scale(1.05); }
    }
    
    .main-header p {
        font-size: 1.4rem;
        margin: 0;
        font-weight: 500;
        color: #FFF8E1;
        text-shadow: 1px 1px 4px rgba(0,0,0,0.5);
        opacity: 0.95;
        letter-spacing: 1px;
    }
    .single-box {
        background: #FFF8E1;
        color: #F57F17;
        padding: 2rem;
        border-radius: 12px;
        border: 3px solid #FFC107;
        margin: 1rem 0;
        box-shadow: 0 8px 20px rgba(255, 193, 7, 0.3);
        text-align: left;
        display: block;
        width: 100%;
    }
    .single-success-box {
        background: #E8F5E8;
        color: #2E7D32;
        padding: 2rem;
        border-radius: 12px;
        border: 3px solid #4CAF50;
        margin: 1rem 0;
        box-shadow: 0 8px 20px rgba(76, 175, 80, 0.3);
        text-align: left;
        display: block;
        width: 100%;
    }
    .warning-box {
        background: #FFF8E1;
        color: #F57F17;
        padding: 2rem;
        border-radius: 12px;
        border: 3px solid #FFC107;
        margin: 1rem 0;
        box-shadow: 0 8px 20px rgba(255, 193, 7, 0.3);
    }
    
    .success-box {
        background: #E8F5E8;
        color: #2E7D32;
        padding: 2rem;
        border-radius: 12px;
        border: 3px solid #4CAF50;
        margin: 1rem 0;
        box-shadow: 0 8px 20px rgba(76, 175, 80, 0.3);
    }
    .stDataFrame {
        max-width: 100%;
        overflow-x: auto;
    }
    .stDataFrame table {
        width: 100%;
        table-layout: auto;
    }
    .stDataFrame th, .stDataFrame td {
        word-wrap: break-word;
        max-width: none;
        overflow: visible;
        text-overflow: clip;
    }
    
    .download-section {
        background: #F5F5F5;
        padding: 2rem;
        border-radius: 15px;
        border: 2px solid #ddd;
        margin: 2rem 0;
        display: block !important;
        visibility: visible !important;
    }
    
    .stButton > button {
        background: linear-gradient(45deg, #FF1744, #E91E63, #9C27B0, #673AB7);
        color: white !important;
        border-radius: 30px;
        border: 3px solid transparent;
        padding: 1rem 2.5rem;
        font-weight: 800;
        font-size: 1.1rem;
        transition: all 0.3s;
        box-shadow: 0 8px 25px rgba(255, 23, 68, 0.4);
        display: block !important;
        visibility: visible !important;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px);
        box-shadow: 0 12px 35px rgba(255, 23, 68, 0.5);
        background: linear-gradient(45deg, #D81B60, #C2185B, #7B1FA2, #512DA8);
        border-color: #FF1744;
    }
    
    .stDownloadButton > button {
        background: linear-gradient(45deg, #2196F3, #03A9F4, #00BCD4, #009688);
        color: white !important;
        border-radius: 15px;
        border: 2px solid #2196F3;
        padding: 0.8rem 1.5rem;
        font-weight: 700;
        font-size: 1rem;
        transition: all 0.3s;
        box-shadow: 0 4px 15px rgba(33, 150, 243, 0.3);
        margin: 0.5rem;
        display: inline-block !important;
        visibility: visible !important;
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(33, 150, 243, 0.4);
        background: linear-gradient(45deg, #1976D2, #0288D1, #0097A7, #00796B);
    }
    
    .element-container {
        display: block !important;
        visibility: visible !important;
    }
    
    .stMarkdown {
        display: block !important;
        visibility: visible !important;
    }
    
    .stExpander {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .metric-container {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid #e9ecef;
        margin: 1rem 0;
    }
    
    .stFileUploader {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .stTextArea {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .stApp .stWidget {
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .highlight-keyword {
        font-weight: bold;
        background-color: #FFF9C4;
        color: #D81B60;
        padding: 2px 4px;
        border-radius: 3px;
    }
</style>
""", unsafe_allow_html=True)

# Function to highlight specific keywords
def highlight_keywords(text):
    keywords = ["Actual", "Revised", "Modified", "Change", "Original"]
    for keyword in keywords:
        text = re.sub(
            rf'\b{keyword}\b',
            f'<span class="highlight-keyword">{keyword}</span>',
            text,
            flags=re.IGNORECASE
        )
    return text

# Initialize AI (Gemini with OpenAI fallback)
@st.cache_data
def initialize_ai(_gemini_api_key, _openai_api_key):
    try:
        if _gemini_api_key:
            genai.configure(api_key=_gemini_api_key)
            model = genai.GenerativeModel('gemini-2.5-flash')
            return model, "gemini"
        elif _openai_api_key:
            client = openai.OpenAI(api_key=_openai_api_key)
            return client, "openai"
        else:
            st.error("❌ Neither Gemini nor OpenAI API key found! Please add one to your .env file or Streamlit Cloud secrets")
            return None, None
    except Exception as e:
        st.error(f"❌ Error initializing AI: {str(e)}")
        return None, None

# Define Pydantic model for comparison table entries
class ComparisonEntry(BaseModel):
    section_element: str = Field(description="The section or element being compared", alias="Section/Element")
    original: str = Field(description="Content in the original document", alias="Original")
    revised: str = Field(description="Content in the revised document", alias="Revised")
    change_type: str = Field(description="Type of change (Added, Removed, Modified, Unchanged)", alias="Change Type")
    where_changed: str = Field(description="Location of the change (e.g., section or paragraph)", alias="Where Changed")
    impact: str = Field(description="Impact of the change", alias="Impact")
    recommendation: str = Field(description="Recommended action", alias="Recommendation")

# Define Pydantic model for the full output
class ComparisonOutput(BaseModel):
    comparison_table: List[ComparisonEntry] = Field(description="List of comparison entries")
    executive_summary: str = Field(description="Concise overview of main changes and impacts")
    section_analysis: str = Field(description="Narrative section-by-section analysis including unchanged content")
    image_comparison: str = Field(description="Image comparison report")
    recommendations_summary: str = Field(description="Consolidated recommendations")

# Initialize the parser
output_parser = PydanticOutputParser(pydantic_object=ComparisonOutput)

# Extract text from PDF with OCR fallback for scanned pages and return page count
def extract_pdf_text(file):
    try:
        file_content = file.read()
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        num_pages = len(pdf_reader.pages)
        text = ""
        has_text_layer = False
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text += page_text + "\n"
                has_text_layer = True

        if not has_text_layer or len(text) < 100:
            st.info("Detected potential scanned PDF. Using OCR for better extraction.")
            pdf_document = fitz.open(stream=io.BytesIO(file_content), filetype="pdf")
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
            pdf_document.close()

        return text.strip(), num_pages
    except Exception as e:
        st.error(f"❌ Error reading PDF: {str(e)}")
        return None, None

# Extract text from Word document and estimate page count
def extract_docx_text(file):
    try:
        doc = docx.Document(io.BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        word_count = len(text.split())
        num_pages = max(1, (word_count // 500) + 1)
        return text.strip(), num_pages
    except Exception as e:
        st.error(f"❌ Error reading Word document: {str(e)}")
        return None, None

# Extract text from JSON document
def extract_json_text(file):
    try:
        data = json.load(io.BytesIO(file.read()))
        def flatten_json(obj, parent_key='', sep='.'):
            items = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    new_key = f"{parent_key}{sep}{k}" if parent_key else k
                    items.extend(flatten_json(v, new_key, sep).items())
            elif isinstance(obj, list):
                for i, v in enumerate(obj):
                    new_key = f"{parent_key}{sep}{i}" if parent_key else str(i)
                    items.extend(flatten_json(v, new_key, sep).items())
            else:
                items.append((parent_key, str(obj)))
            return dict(items)
        
        flat_data = flatten_json(data)
        text = "\n".join([f"{key}: {value}" for key, value in flat_data.items()])
        line_count = len(text.splitlines())
        num_pages = max(1, (line_count // 50) + 1)
        return text.strip(), num_pages
    except Exception as e:
        st.error(f"❌ Error reading JSON: {str(e)}")
        return None, None

# Extract text based on file type
def extract_text(file):
    file.seek(0)
    if file.type == "application/pdf":
        return extract_pdf_text(file)
    elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        return extract_docx_text(file)
    elif file.type == "text/plain":
        text = str(file.read(), "utf-8")
        line_count = len(text.splitlines())
        num_pages = max(1, (line_count // 50) + 1)
        return text.strip(), num_pages
    elif file.type == "application/json":
        return extract_json_text(file)
    else:
        st.error("❌ Unsupported file format! Please upload PDF, Word, TXT, or JSON files.")
        return None, None

# Validate the comparison table
def validate_comparison_table(table_data):
    expected_columns = ["Section/Element", "Original", "Revised", "Change Type", "Where Changed", "Impact", "Recommendation"]
    valid_change_types = ["Added", "Removed", "Modified", "Unchanged"]
    
    if not table_data:
        return False, "Comparison table is empty"
    
    for row in table_data:
        if not all(col in row for col in expected_columns):
            return False, f"Row missing required columns: {row}"
        if row["Change Type"] not in valid_change_types:
            return False, f"Invalid Change Type: {row['Change Type']}"
        for col in expected_columns:
            if col == "Revised" and row["Change Type"] == "Removed" and not row[col]:
                continue
            if col == "Original" and row["Change Type"] == "Added" and not row[col]:
                continue
            if not row[col]:
                return False, f"Row contains empty required field '{col}': {row}"
    
    return True, "Comparison table is valid"

# Normalize change type
def normalize_change_type(change_type):
    valid_types = ["Added", "Removed", "Modified", "Unchanged"]
    if not isinstance(change_type, str):
        return "Modified"  # Default for non-string inputs
    change_type = change_type.strip()
    
    # Check for exact match
    if change_type in valid_types:
        return change_type
    
    # Check for substring match (case-insensitive)
    change_type_lower = change_type.lower()
    for valid_type in valid_types:
        if valid_type.lower() in change_type_lower:
            return valid_type
            
    # Default to Modified for any unrecognized change type
    return "Modified"

# Compare documents using AI with LangChain parser
def compare_documents_ai(original_text, revised_text, model, api_type):
    try:
        prompt_template = PromptTemplate(
            input_variables=["original_text", "revised_text"],
            template="""
You are an expert AI assistant specialized in document comparison and impact analysis.

Your task is to compare two documents:
1. Original Document – the first version.
2. Revised Document – the updated version, where changes may exist.

The documents may be in PDF, DOCX, TXT, or JSON formats, and may include tables, lists, or images.

**Comparative Analysis Guidelines**
1. **Section-Level & Criteria-Level Comparison**
   - Compare at the criteria, section, and paragraph level, not word by word.
   - Classify each part as:
     - Added: New in the revised version.
     - Removed: Present in the original but missing in the revised version.
     - Modified: Present in both but with meaningful changes.
     - Unchanged: Identical content preserved.
   - Specify where the change occurred (section, heading, paragraph reference).

2. **Detailed Reporting**
   - List what changed, where it changed, and how it changed.
   - For modifications, show Original → Revised.
   - For additions, use "(Empty)" for the Original field if no content existed.
   - For deletions, use "(Empty)" for the Revised field if content was removed.
   - Include unchanged content for context.

3. **Image Comparison**
   - Detect if images were added, removed, or modified.
   - Report captions, labels, or embedded text.
   - If modified, explain what was altered (e.g., updated chart values, new diagram element).
   - Provide image references (page number, section, or filename if available).

4. **Impact Analysis**
   - For each change, explain the nature and assess the impact (e.g., stricter eligibility, expanded access, added compliance).
   - Flag critical updates with significant operational or compliance implications.

5. **Recommendations**
   - For Additions: Suggest how to integrate the new requirement (e.g., update training, notify workers, revise compliance checklists).
   - For Removals: Suggest whether old requirements need to be archived or communicated as obsolete.
   - For Modifications: Suggest what policies, workflows, or worker groups are affected and what adjustments are required.
   - For Unchanged Content: Confirm no action is required but note continued compliance.

**Output Format**
Provide the output as JSON with the following structure:
{format_instructions}

The output must include a comparison table with the following columns:
| Section/Element | Original | Revised | Change Type | Where Changed | Impact | Recommendation |
|-----------------|----------|---------|-------------|---------------|--------|---------------|
| (Example) Eligibility Criteria | "2 years experience" | "3 years experience" | Modified | Section 2.1 | Stricter eligibility | Notify HR to update hiring requirements |
| (Example) Safety Diagram | Present | (Empty) | Removed | Page 4, Figure A | Loss of visual compliance guide | Archive old diagram and ensure staff are informed |
| (Example) Training Requirement | (Empty) | "Safety training required" | Added | Section 3.2 | New compliance rule | Schedule mandatory training sessions |
| (Example) Age Requirement | 18+ | 18+ | Unchanged | Section 1.1 | No impact | No action needed |

Ensure the comparison_table includes ALL columns: Section/Element, Original, Revised, Change Type, Where Changed, Impact, Recommendation. Every row must have all columns populated, using "(Empty)" for missing content in Original or Revised fields as appropriate.

Followed by:
3. Narrative Section-by-Section Analysis – including unchanged content.
4. Image Comparison Report – added/removed/modified images with details.
5. Recommendations Summary – consolidated, clear next steps for the client/admin.

**ACTUAL DOCUMENT**:
{original_text}

**REVISED DOCUMENT**:
{revised_text}
            """,
            partial_variables={"format_instructions": output_parser.get_format_instructions()}
        )
        
        prompt = prompt_template.format(
            original_text=original_text,
            revised_text=revised_text
        )
        
        placeholder = st.empty()
        text = ""
        chunk_count = 0
        
        if api_type == "gemini":
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(temperature=0.1, top_p=0.9),
                stream=True
            )
            for chunk in response:
                if hasattr(chunk, 'text'):
                    text += chunk.text
                chunk_count += 1
                if chunk_count % 5 == 0:
                    placeholder.text(f"Processing... received {chunk_count} chunks")
        else:  # OpenAI
            response = model.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                top_p=0.9,
                stream=True
            )
            for chunk in response:
                if chunk.choices[0].delta.content:
                    text += chunk.choices[0].delta.content
                chunk_count += 1
                if chunk_count % 5 == 0:
                    placeholder.text(f"Processing... received {chunk_count} chunks")
        
        placeholder.text("Parsing response...")
        parsed_output = output_parser.parse(text)
        placeholder.empty()
        
        parsed_output_dict = parsed_output.dict(by_alias=True)
        
        for row in parsed_output_dict["comparison_table"]:
            for key in ["Section/Element", "Original", "Revised", "Change Type", "Where Changed", "Impact", "Recommendation"]:
                if key not in row:
                    row[key] = "(Empty)"
                else:
                    if key == "Change Type":
                        row[key] = normalize_change_type(row[key])
                    else:
                        row[key] = clean_br_tags(row[key])
        
        is_valid, validation_message = validate_comparison_table(parsed_output_dict["comparison_table"])
        if not is_valid:
            st.warning(f"⚠️ Validation failed: {validation_message}")
        
        return parsed_output_dict
    except Exception as e:
        st.error(f"❌ Error in AI comparison: {str(e)}")
        return None

# Clean <br> tags from text
def clean_br_tags(text):
    if isinstance(text, str):
        return re.sub(r'<br\s*/?>', ' ', text, flags=re.IGNORECASE)
    return text

# Generate comparison analysis table (high-level metrics)
def generate_comparison_table(original_text, revised_text, original_filename, revised_filename):
    similarity = calculate_similarity(original_text, revised_text)
    length_change = ((len(revised_text) - len(original_text)) / len(original_text) * 100) if len(original_text) > 0 else 0
    line_diff = len(revised_text.splitlines()) - len(original_text.splitlines())
    word_diff = len(revised_text.split()) - len(original_text.split())

    original_filename = clean_br_tags(original_filename[:20] + "..." if len(original_filename) > 20 else original_filename)
    revised_filename = clean_br_tags(revised_filename[:20] + "..." if len(revised_filename) > 20 else revised_filename)

    table_data = [
        {
            "Metric": "Document Names",
            "Original": original_filename,
            "Revised": revised_filename,
            "Description": "Compared document names"
        },
        {
            "Metric": "Length Change",
            "Original": f"{len(original_text):,} chars",
            "Revised": f"{len(revised_text):,} chars",
            "Description": f"{length_change:+.1f}% char change"
        },
        {
            "Metric": "Line Difference",
            "Original": f"{len(original_text.splitlines()):,} lines",
            "Revised": f"{len(revised_text.splitlines()):,} lines",
            "Description": f"{line_diff:+,} lines changed"
        },
        {
            "Metric": "Word Difference",
            "Original": f"{len(original_text.split()):,} words",
            "Revised": f"{len(revised_text.split()):,} words",
            "Description": f"{word_diff:+,} words changed"
        },
    ]
    for row in table_data:
        for key in row:
            row[key] = clean_br_tags(row[key])
    return table_data

def generate_pdf_report(ai_analysis, original_filename, revised_filename, metrics_table, comparison_table):
    try:
        ai_analysis = {k: clean_br_tags(v) if isinstance(v, str) else v for k, v in ai_analysis.items()}
        original_filename = clean_br_tags(original_filename)
        revised_filename = clean_br_tags(revised_filename)
        metrics_table = [{k: clean_br_tags(v) for k, v in row.items()} for row in metrics_table]
        comparison_table = [{k: clean_br_tags(v) if v else "(Empty)" for k, v in row.items()} for row in comparison_table]

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(letter),
            rightMargin=30,
            leftMargin=30,
            topMargin=30,
            bottomMargin=50
        )
        elements = []

        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading_style = styles['Heading1']
        body_style = styles['Normal']
        body_style.fontSize = 7
        body_style.leading = 9
        body_style.wordWrap = 'CJK'

        elements.append(Paragraph("Document Comparison Report", title_style))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Original Document: {original_filename}", body_style))
        elements.append(Paragraph(f"Revised Document: {revised_filename}", body_style))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body_style))
        elements.append(Spacer(1, 18))

        elements.append(Paragraph("Comparison Metrics Table", heading_style))
        metrics_data = [["Metric", "Original", "Revised", "Description"]] + [
            [Paragraph(row["Metric"], body_style), 
             Paragraph(row["Original"], body_style), 
             Paragraph(row["Revised"], body_style), 
             Paragraph(row["Description"], body_style)]
            for row in metrics_table
        ]
        metrics_col_widths = [100, 120, 120, 200]
        metrics_table_obj = Table(metrics_data, colWidths=metrics_col_widths, repeatRows=1, splitByRow=1)
        metrics_table_obj.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 7),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            ('LEFTPADDING', (0, 0), (-1, -1), 2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('WORDWRAP', (0, 0), (-1, -1), 1),
        ]))
        elements.append(metrics_table_obj)
        elements.append(Spacer(1, 18))

        elements.append(Paragraph("AI Comparison Table", heading_style))
        comparison_data = [["Section/Element", "Original", "Revised", "Change Type", "Where Changed", "Impact", "Recommendation"]] + [
            [
                Paragraph(row.get("Section/Element", "")[:100], body_style),
                Paragraph(row.get("Original", "")[:200], body_style),
                Paragraph(row.get("Revised", "")[:200], body_style),
                Paragraph(row.get("Change Type", ""), body_style),
                Paragraph(row.get("Where Changed", ""), body_style),
                Paragraph(row.get("Impact", "")[:150], body_style),
                Paragraph(row.get("Recommendation", "")[:150], body_style)
            ]
            for row in comparison_table
        ]
        comparison_col_widths = [80, 100, 100, 50, 50, 80, 100]
        comparison_table_obj = Table(comparison_data, colWidths=comparison_col_widths, repeatRows=1, splitByRow=1)
        comparison_table_obj.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 6),
            ('FONTSIZE', (0, 1), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 1),
            ('RIGHTPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADING', (0, 0), (-1, -1), 1),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('WORDWRAP', (0, 0), (-1, -1), 1),
        ]))
        elements.append(comparison_table_obj)
        elements.append(Spacer(1, 18))

        elements.append(Paragraph("Executive Summary", heading_style))
        elements.append(Paragraph(ai_analysis["executive_summary"], body_style))
        elements.append(Spacer(1, 12))

        elements.append(Paragraph("Section-by-Section Analysis", heading_style))
        elements.append(Paragraph(ai_analysis["section_analysis"], body_style))
        elements.append(Spacer(1, 12))

        elements.append(Paragraph("Image Comparison Report", heading_style))
        elements.append(Paragraph(ai_analysis["image_comparison"], body_style))
        elements.append(Spacer(1, 12))

        elements.append(Paragraph("Recommendations Summary", heading_style))
        elements.append(Paragraph(ai_analysis["recommendations_summary"], body_style))
        elements.append(Spacer(1, 12))

        doc.build(elements)
        buffer.seek(0)
        return buffer

    except Exception as e:
        st.error(f"❌ Error generating PDF: {str(e)}")
        return None

# Generate DOCX report
def generate_docx_report(ai_analysis, original_filename, revised_filename, metrics_table, comparison_table):
    try:
        ai_analysis = {k: clean_br_tags(v) if isinstance(v, str) else v for k, v in ai_analysis.items()}
        original_filename = clean_br_tags(original_filename)
        revised_filename = clean_br_tags(revised_filename)
        metrics_table = [{k: clean_br_tags(v) for k, v in row.items()} for row in metrics_table]
        comparison_table = [{k: clean_br_tags(v) if v else "(Empty)" for k, v in row.items()} for row in comparison_table]

        doc = Document()
        title = doc.add_heading('Document Comparison Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Original Document: {original_filename}")
        doc.add_paragraph(f"Revised Document: {revised_filename}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading('Comparison Metrics Table', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Metric", "Original", "Revised", "Description"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        col_widths = [3.0, 4.0, 4.0, 3.0]
        for col_idx, width in enumerate(col_widths):
            table.columns[col_idx].width = Cm(width)
        for row in metrics_table:
            row_cells = table.add_row().cells
            row_cells[0].text = row["Metric"][:100]
            row_cells[1].text = row["Original"][:150]
            row_cells[2].text = row["Revised"][:150]
            row_cells[3].text = row["Description"][:150]

        doc.add_heading('AI Comparison Table', level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Section/Element", "Original", "Revised", "Change Type", "Where Changed", "Impact", "Recommendation"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        col_widths = [2.0, 2.5, 2.5, 1.5, 1.5, 2.0, 2.5]
        for col_idx, width in enumerate(col_widths):
            table.columns[col_idx].width = Cm(width)
        for row in comparison_table:
            row_cells = table.add_row().cells
            for i, col in enumerate(headers):
                row_cells[i].text = row.get(col, "")[:150]
                for paragraph in row_cells[i].paragraphs:
                    paragraph.style.font.size = Cm(0.25)

        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(ai_analysis["executive_summary"])
        doc.add_heading('Section-by-Section Analysis', level=1)
        doc.add_paragraph(ai_analysis["section_analysis"])
        doc.add_heading('Image Comparison Report', level=1)
        doc.add_paragraph(ai_analysis["image_comparison"])
        doc.add_heading('Recommendations Summary', level=1)
        doc.add_paragraph(ai_analysis["recommendations_summary"])

        footer = doc.add_paragraph('AI Document Comparator')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Error generating DOCX: {str(e)}")
        return None

# Generate TXT report
def generate_txt_report(ai_analysis, original_filename, revised_filename, metrics_table, comparison_table, similarity, length_change, original_text, revised_text):
    try:
        ai_analysis = {k: clean_br_tags(v) if isinstance(v, str) else v for k, v in ai_analysis.items()}
        original_filename = clean_br_tags(original_filename)
        revised_filename = clean_br_tags(revised_filename)
        metrics_table = [{k: clean_br_tags(v) for k, v in row.items()} for row in metrics_table]
        comparison_table = [{k: clean_br_tags(v) if v else "(Empty)" for k, v in row.items()} for row in comparison_table]

        table_text = "COMPARISON METRICS TABLE:\n"
        table_text += f"{'Metric':<15} {'Original':<20} {'Revised':<20} {'Description':<25}\n"
        table_text += "-" * 80 + "\n"
        for row in metrics_table:
            table_text += f"{row['Metric'][:15]:<15} {row['Original'][:20]:<20} {row['Revised'][:20]:<20} {row['Description'][:25]:<25}\n"

        ai_table_text = "\nAI COMPARISON TABLE:\n"
        headers = ["Section/Element", "Original", "Revised", "Change Type", "Where Changed", "Impact", "Recommendation"]
        header_str = "  ".join([f"{h:<30}" for h in headers])
        ai_table_text += header_str + "\n"
        ai_table_text += "-" * len(header_str) + "\n"
        for row in comparison_table:
            row_data = [row.get(col, "")[:30] for col in headers]
            ai_table_text += "  ".join([f"{c:<30}" for c in row_data]) + "\n"

        report_text = f"""Document Comparison Report

Original Document: {original_filename}
Revised Document: {revised_filename}
Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Generated By: AI Document Comparator

COMPARISON METRICS:
- Similarity: {similarity}%
- Length Change: {length_change:+.1f}%
- Character Difference: {len(revised_text) - len(original_text):+,}
- Line Difference: {len(revised_text.splitlines()) - len(original_text.splitlines()):+,}

{table_text}
{ai_table_text}
EXECUTIVE SUMMARY:
{ai_analysis["executive_summary"]}

SECTION-BY-SECTION ANALYSIS:
{ai_analysis["section_analysis"]}

IMAGE COMPARISON REPORT:
{ai_analysis["image_comparison"]}

RECOMMENDATIONS SUMMARY:
{ai_analysis["recommendations_summary"]}
"""
        buffer = BytesIO()
        buffer.write(report_text.encode('utf-8'))
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Error generating TXT: {str(e)}")
        return None

# Generate JSON report
def generate_json_report(ai_analysis, original_filename, revised_filename, metrics_table, comparison_table):
    try:
        ai_analysis = {k: clean_br_tags(v) if isinstance(v, str) else v for k, v in ai_analysis.items()}
        original_filename = clean_br_tags(original_filename)
        revised_filename = clean_br_tags(revised_filename)
        metrics_table = [{k: clean_br_tags(v) for k, v in row.items()} for row in metrics_table]
        comparison_table = [{k: clean_br_tags(v) if v else "(Empty)" for k, v in row.items()} for row in comparison_table]

        report_data = {
            "report_metadata": {
                "original_document": original_filename,
                "revised_document": revised_filename,
                "analysis_date": datetime.now().isoformat(),
                "generated_by": "AI Document Comparator"
            },
            "comparison_metrics_table": metrics_table,
            "ai_comparison_table": comparison_table,
            "executive_summary": ai_analysis["executive_summary"],
            "section_analysis": ai_analysis["section_analysis"],
            "image_comparison": ai_analysis["image_comparison"],
            "recommendations_summary": ai_analysis["recommendations_summary"],
            "document_stats": {
                "analysis_type": "AI-powered document comparison",
                "comparison_method": "AI analysis with LangChain parser"
            }
        }
        buffer = BytesIO()
        buffer.write(json.dumps(report_data, indent=4, ensure_ascii=False).encode('utf-8'))
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"❌ Error generating JSON: {str(e)}")
        return None

# Calculate similarity percentage
def calculate_similarity(text1, text2):
    try:
        matcher = difflib.SequenceMatcher(None, text1, text2)
        return round(matcher.ratio() * 100, 1)
    except:
        return 0.0

# Main application
def main():
    # Initialize session state variables
    if 'metrics_table' not in st.session_state:
        st.session_state['metrics_table'] = None
    if 'ai_analysis' not in st.session_state:
        st.session_state['ai_analysis'] = None
    if 'analysis_date' not in st.session_state:
        st.session_state['analysis_date'] = None

    st.markdown("""
    <div class="main-header">
        <h1>🤖 AI Document Comparator</h1>
        <p>Upload any two documents and let AI analyze the differences for you!</p>
    </div>
    """, unsafe_allow_html=True)

    # Load API keys
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not gemini_api_key and not openai_api_key:
        st.error("❌ Neither GEMINI_API_KEY nor OPENAI_API_KEY found. Please set one in Streamlit Cloud secrets or .env file.")
        st.stop()

    # Initialize AI model
    model, api_type = initialize_ai(gemini_api_key, openai_api_key)
    if not model:
        st.stop()

    with st.sidebar:
        st.markdown("### 🔧 System Status")
        if st.button("🔍 Test API Connection"):
            if gemini_api_key or openai_api_key:
                st.success(f"✅ {'Gemini' if gemini_api_key else 'OpenAI'} API key found")
                try:
                    if gemini_api_key:
                        genai.configure(api_key=gemini_api_key)
                        test_model = genai.GenerativeModel('gemini-2.5-flash')
                        test_response = test_model.generate_content("Test connection.")
                    else:
                        test_client = openai.OpenAI(api_key=openai_api_key)
                        test_response = test_client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{"role": "user", "content": "Test connection."}]
                        )
                    st.success("✅ API connection successful!")
                except Exception as e:
                    st.error(f"❌ API connection failed: {str(e)}")
            else:
                st.error("❌ No API key found")
        
        st.markdown("### 📋 How to Use")
        st.markdown("""
        1. Upload Original Document - First version/baseline document
        2. Upload Revised Document - Updated/modified version
        3. Click Compare - AI analyzes all differences
        4. Review Results - Get comprehensive comparison report
        5. Download Report - Save results in multiple formats
        """)
        
        st.markdown("### 📄 Supported Formats")
        st.markdown("- PDF (.pdf) - All types of PDF documents (with OCR for scanned)")
        st.markdown("- Word (.docx) - Microsoft Word documents")
        st.markdown("- Text (.txt) - Plain text files")
        st.markdown("- JSON (.json) - JSON structured data")
        
        st.markdown("### ⚡ Key Features")
        st.markdown("- AI-Powered Analysis - Advanced comparison using Gemini or OpenAI with LangChain parsing")
        st.markdown("- Comprehensive Reports - Detailed change analysis with validated comparison table")
        st.markdown("- Multiple Export Formats - PDF, DOCX, TXT, JSON")
        st.markdown("- Visual Diff View - Technical section-by-section comparison")
        st.markdown("- Document Preview - Table Comparison view")
        st.markdown("- OCR Support - Automatic OCR for scanned PDFs without text layer")

        st.markdown("### ℹ️ Note")
        st.markdown("AI outputs may vary slightly due to the generative nature of the model, but LangChain parsing ensures consistent table structure.")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.subheader("📄 Original Document")
        original_file = st.file_uploader(
            "Choose the original/baseline document",
            type=['pdf', 'docx', 'txt', 'json'],
            key="original",
            help="Upload the first version of your document for comparison"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.subheader("📄 Revised Document") 
        revised_file = st.file_uploader(
            "Choose the revised/updated document",
            type=['pdf', 'docx', 'txt', 'json'],
            key="revised",
            help="Upload the modified version of your document"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    if original_file and revised_file:
        with st.spinner("🔍 Extracting text from documents..."):
            original_text, original_pages = extract_text(original_file)
            revised_text, revised_pages = extract_text(revised_file)

        if original_text and revised_text:
            similarity = calculate_similarity(original_text, revised_text)
            length_change = ((len(revised_text) - len(original_text)) / len(original_text) * 100) if len(original_text) > 0 else 0
            char_diff = len(revised_text) - len(original_text)
            line_diff = len(revised_text.splitlines()) - len(original_text.splitlines())
            word_diff = len(revised_text.split()) - len(original_text.split())
            
            col1, col2 = st.columns(2)
            with col1:
                pages_str = f"Pages: {original_pages:,}<br>" if original_pages else ""
                st.markdown(f"""
                <div class="single-success-box">
                    <h4>📄 Original Document</h4>
                    File: {clean_br_tags(original_file.name)}<br>
                    Size: {len(original_text):,} characters<br>
                    Lines: {len(original_text.splitlines()):,}<br>
                    Words: {len(original_text.split()):,}<br>
                    {pages_str}
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                pages_str = f"Pages: {revised_pages:,}<br>" if revised_pages else ""
                st.markdown(f"""
                <div class="single-success-box">
                    <h4>📄 Revised Document</h4>
                    File: {clean_br_tags(revised_file.name)}<br>
                    Size: {len(revised_text):,} characters<br>
                    Lines: {len(revised_text.splitlines()):,}<br>
                    Words: {len(revised_text.split()):,}<br>
                    {pages_str}
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown(f"""
            <div class="metric-container">
                <h4>📊 Quick Metrics</h4>
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 1rem;">
                    <div>Similarity: {similarity}%</div>
                    <div>Length Change: {length_change:+.1f}%</div>
                    <div>Size Difference: {char_diff:+,} chars</div>
                    <div>Line Difference: {line_diff:+,}</div>
                    <div>Word Difference: {word_diff:+,}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🔍 COMPARE DOCUMENTS", type="primary", use_container_width=True):
                with st.spinner("🤖 AI is analyzing the differences... This may take a moment."):
                    metrics_table = generate_comparison_table(
                        original_text,
                        revised_text,
                        original_file.name,
                        revised_file.name
                    )
                    st.session_state['metrics_table'] = metrics_table
                    
                    ai_analysis = compare_documents_ai(original_text, revised_text, model, api_type)
                    if ai_analysis:
                        if not ai_analysis['comparison_table']:
                            ai_analysis['comparison_table'] = [{
                                "Section/Element": "No changes",
                                "Original": "(Empty)",
                                "Revised": "(Empty)",
                                "Change Type": "Unchanged",
                                "Where Changed": "N/A",
                                "Impact": "No impact",
                                "Recommendation": "No action needed"
                            }]
                        st.session_state['ai_analysis'] = ai_analysis
                        st.session_state['analysis_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            if st.session_state['ai_analysis'] and st.session_state['metrics_table']:
                st.markdown('<div class="comparison-result">', unsafe_allow_html=True)
                st.markdown("## 📊 Comparison Metrics Table")
                st.dataframe(st.session_state['metrics_table'], use_container_width=True)
                
                st.markdown("## 📊 AI Comparison Table")
                df = pd.DataFrame(st.session_state['ai_analysis']['comparison_table'])
                columns = ["Section/Element", "Original", "Revised", "Change Type", "Where Changed", "Impact", "Recommendation"]
                for col in columns:
                    if col not in df.columns:
                        df[col] = "(Empty)"
                df = df[columns]
                st.dataframe(df, use_container_width=True)
                
                st.markdown("## 🤖 AI Analysis Results")
                st.markdown('<div class="difference-box">', unsafe_allow_html=True)
                st.markdown(f"**Executive Summary**\n{highlight_keywords(clean_br_tags(st.session_state['ai_analysis']['executive_summary']))}", unsafe_allow_html=True)
                st.markdown(f"**Section-by-Section Analysis**\n{highlight_keywords(clean_br_tags(st.session_state['ai_analysis']['section_analysis']))}", unsafe_allow_html=True)
                st.markdown(f"**Image Comparison Report**\n{highlight_keywords(clean_br_tags(st.session_state['ai_analysis']['image_comparison']))}", unsafe_allow_html=True)
                st.markdown(f"**Recommendations Summary**\n{highlight_keywords(clean_br_tags(st.session_state['ai_analysis']['recommendations_summary']))}", unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("## 📖 Document Preview")
            preview_col1, preview_col2 = st.columns(2)
            with preview_col1:
                st.markdown("### Original Document")
                st.text_area(
                    "Original Content",
                    highlight_keywords(clean_br_tags(original_text)),
                    height=400,
                    disabled=False,
                    key="original_preview"
                )
            
            with preview_col2:
                st.markdown("### Revised Document") 
                st.text_area(
                    "Revised Content",
                    highlight_keywords(clean_br_tags(revised_text)),
                    height=400,
                    disabled=False,
                    key="revised_preview"
                )
            
            if st.session_state['ai_analysis'] and st.session_state['metrics_table']:
                st.markdown("## 💾 Download Analysis Report")
                metrics_table = st.session_state['metrics_table']
                ai_analysis = st.session_state['ai_analysis']
                comparison_table = ai_analysis['comparison_table']
                analysis_date = st.session_state.get('analysis_date', 'N/A')
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    txt_buffer = generate_txt_report(
                        ai_analysis,
                        original_file.name,
                        revised_file.name,
                        metrics_table,
                        comparison_table,
                        similarity,
                        length_change,
                        original_text,
                        revised_text
                    )
                    if txt_buffer:
                        st.download_button(
                            label="📄 Download TXT",
                            data=txt_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                
                with col2:
                    pdf_buffer = generate_pdf_report(
                        ai_analysis,
                        original_file.name,
                        revised_file.name,
                        metrics_table,
                        comparison_table
                    )
                    if pdf_buffer:
                        st.download_button(
                            label="📄 Download PDF",
                            data=pdf_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                
                with col3:
                    docx_buffer = generate_docx_report(
                        ai_analysis,
                        original_file.name,
                        revised_file.name,
                        metrics_table,
                        comparison_table
                    )
                    if docx_buffer:
                        st.download_button(
                            label="📄 Download DOCX",
                            data=docx_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                
                with col4:
                    json_buffer = generate_json_report(
                        ai_analysis,
                        original_file.name,
                        revised_file.name,
                        metrics_table,
                        comparison_table
                    )
                    if json_buffer:
                        st.download_button(
                            label="📄 Download JSON",
                            data=json_buffer,
                            file_name=f"comparison_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                            mime="application/json",
                            use_container_width=True
                        )

    else:
        st.markdown("""
        <div class="single-box">
            <h4>⏳ Ready to Compare Documents</h4>
            Please upload both documents (original and revised versions) to begin the AI-powered comparison analysis.<br>
            Supported file types: PDF, DOCX, TXT, JSON<br>
            Analysis includes: Content changes, structural differences, impact assessment, validated comparison table, and detailed recommendations.
        </div>
        """, unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align: center; color: #666; margin-top: 2rem; padding: 1rem;">
        <p><strong>🤖 AI Document Comparator</strong> | Powered by Google Gemini or OpenAI and LangChain | Built with Streamlit</p>
        <p>Compare any documents - contracts, reports, articles, code, essays, manuals, JSON data, and more!</p>
        <p><em>Universal document comparison for all your analysis needs</em></p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
